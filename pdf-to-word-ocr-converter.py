import os
import logging
from pdf2image import convert_from_path
from pdf2image.exceptions import PDFInfoNotInstalledError
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import cv2
import numpy as np
import traceback
import sys

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ConversionError(Exception):
    """Custom exception for conversion errors"""
    pass

def preprocess_image(image):
    logger.debug("Preprocessing image")
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
        # Apply threshold to get image with only black and white
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
        return Image.fromarray(thresh)
    except Exception as e:
        logger.error(f"Error in preprocessing image: {str(e)}")
        raise ConversionError("Failed to preprocess image") from e

def perform_ocr(image, config):
    logger.debug("Performing OCR")
    try:
        return pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config=config)
    except Exception as e:
        logger.error(f"Error in OCR: {str(e)}")
        raise ConversionError("OCR failed") from e

def convert_pdf_to_word(pdf_filename, output_filename):
    # Construct full paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    pdf_path = os.path.join(script_dir, 'pdf', pdf_filename)
    output_path = os.path.join(script_dir, 'output', output_filename)
    
    logger.info(f"Converting PDF: {pdf_path}")
    logger.info(f"Output will be saved to: {output_path}")

    try:
        # Check if input file exists
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Input PDF file not found: {pdf_path}")

        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Convert PDF to images
        logger.debug("Converting PDF to images")
        images = convert_from_path(pdf_path)
        
        # Create a new Word document
        doc = Document()
        
        for i, image in enumerate(images):
            logger.info(f"Processing page {i+1}/{len(images)}")
            
            # Preprocess the image
            processed_image = preprocess_image(image)
            
            # Perform OCR with custom configuration
            custom_config = r'--oem 3 --psm 6 -c textord_old_xheight=1'
            ocr_data = perform_ocr(processed_image, custom_config)
            
            # Process OCR data
            logger.debug(f"Processing OCR data for page {i+1}")
            for j in range(len(ocr_data['text'])):
                if int(ocr_data['conf'][j]) > 60:  # Only consider text with confidence > 60%
                    # Add a new paragraph for each text element
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(ocr_data['text'][j])
                    
                    # Set font properties
                    font = run.font
                    font.name = 'Arial'  # You can change this to match the detected font
                    font.size = Pt(12)  # Adjust size as needed
                    
                    # Set alignment based on the position of the text
                    if ocr_data['left'][j] < processed_image.width / 3:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif ocr_data['left'][j] > 2 * processed_image.width / 3:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    logger.debug(f"Added text: '{ocr_data['text'][j]}' with confidence {ocr_data['conf'][j]}")
            
            # Add a page break after each page except the last one
            if i < len(images) - 1:
                doc.add_page_break()
        
        # Save the document
        logger.info(f"Saving document to {output_path}")
        doc.save(output_path)
        logger.info("Conversion completed successfully")

    except FileNotFoundError as e:
        logger.error(f"File not found: {str(e)}")
        raise
    except PDFInfoNotInstalledError:
        logger.error("Poppler is not installed or not in PATH")
        print("Error: Poppler is not installed or not in PATH. Please install Poppler and add it to your system's PATH.")
        print("Installation instructions:")
        if sys.platform.startswith('win'):
            print("1. Download Poppler for Windows: http://blog.alivate.com.au/poppler-windows/")
            print("2. Extract the downloaded file and add the 'bin' folder to your system's PATH")
        elif sys.platform.startswith('darwin'):
            print("1. Install Homebrew if not already installed: https://brew.sh/")
            print("2. Run: brew install poppler")
        else:
            print("1. Run: sudo apt-get install poppler-utils")
        sys.exit(1)
    except ConversionError as e:
        logger.error(f"Conversion error: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        logger.debug(f"Traceback: {traceback.format_exc()}")
        raise ConversionError("An unexpected error occurred during conversion") from e

# Usage
if __name__ == "__main__":
    pdf_filename = 'input.pdf'
    output_filename = 'output.docx'
    try:
        convert_pdf_to_word(pdf_filename, output_filename)
    except Exception as e:
        logger.error(f"Conversion failed: {str(e)}")
        print(f"Error: {str(e)}")
    else:
        print("Conversion completed successfully!")